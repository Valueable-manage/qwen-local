"""
RAG 强化检索流程：
  1. parse_docx()     — 解析 docx，提取问答对
  2. init_rag()       — 向量化存入 ChromaDB（含 BM25 索引）
  3. search()         — 混合检索：向量召回 + BM25 关键词召回 + 重排序
  4. rag_chat()       — 注入 top-3 上下文后回答

强化点：
  ① 混合检索：向量（语义）+ BM25（关键词）双路召回，取并集
  ② 查询改写：用模型把口语化问题扩展为专业术语，提升召回率
  ③ 交叉编码器重排序：用 cross-encoder 对候选结果精排，保留最相关 top-3
  ④ 相似度过滤：过滤掉相关度过低的结果，避免注入无关内容

依赖安装：
  uv pip install python-docx sentence-transformers chromadb jieba rank-bm25
"""

import json
import os
import re
import pickle

# ──────────────────────────────────────────
# 配置
# ──────────────────────────────────────────

CHROMA_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "rag_db", "chroma")
BM25_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "rag_db", "bm25.pkl")
ENTRIES_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "rag_db", "entries.json")
COLLECTION_NAME = "ship_qa"
EMBED_MODEL = "BAAI/bge-small-zh-v1.5"  # 向量嵌入模型
RERANK_MODEL = "BAAI/bge-reranker-base"  # 交叉编码器重排序模型（约 280MB）
SIMILARITY_THRESHOLD = 0.45  # 低于此分数的结果过滤掉（提高以减少无关注入）

os.makedirs(os.path.dirname(CHROMA_DIR), exist_ok=True)

# ──────────────────────────────────────────
# 懒加载单例
# ──────────────────────────────────────────

_embedder = None
_reranker = None
_collection = None
_bm25 = None
_bm25_docs = []  # BM25 对应的原始文本列表
_bm25_metas = []  # BM25 对应的 metadata 列表


def _get_embedder():
    global _embedder
    if _embedder is None:
        from sentence_transformers import SentenceTransformer

        print(f"  加载嵌入模型 {EMBED_MODEL}...")
        _embedder = SentenceTransformer(EMBED_MODEL)
    return _embedder


def _get_reranker():
    global _reranker
    if _reranker is None:
        try:
            from sentence_transformers import CrossEncoder

            print(f"  加载重排序模型 {RERANK_MODEL}...")
            _reranker = CrossEncoder(RERANK_MODEL, max_length=512)
        except Exception as e:
            print(f"  重排序模型加载失败（{e}），跳过重排序")
            _reranker = None
    return _reranker


def _get_collection():
    global _collection
    if _collection is None:
        import chromadb

        client = chromadb.PersistentClient(path=CHROMA_DIR)
        _collection = client.get_or_create_collection(
            name=COLLECTION_NAME, metadata={"hnsw:space": "cosine"}
        )
    return _collection


def _get_bm25():
    global _bm25, _bm25_docs, _bm25_metas
    if _bm25 is None and os.path.exists(BM25_PATH):
        with open(BM25_PATH, "rb") as f:
            saved = pickle.load(f)
        _bm25 = saved["bm25"]
        _bm25_docs = saved["docs"]
        _bm25_metas = saved["metas"]
    return _bm25


# ──────────────────────────────────────────
# 1. 读取文档文本
# ──────────────────────────────────────────


def _read_doc_text(path: str) -> str:
    """从 .doc 或 .docx 提取纯文本"""
    path = os.path.abspath(path)
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        from docx import Document

        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    if ext == ".doc":
        try:
            import win32com.client

            w = win32com.client.Dispatch("Word.Application")
            w.Visible = False
            doc = w.Documents.Open(path)
            text = doc.Content.Text
            doc.Close(False)
            w.Quit()
            if isinstance(text, bytes):
                for enc in ("utf-8", "gbk", "gb2312", "cp936"):
                    try:
                        text = text.decode(enc)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    text = text.decode("utf-8", errors="replace")
            return text or ""
        except Exception as e:
            import warnings
            warnings.warn(f"读取 .doc 失败，已跳过该文件（{path}）: {e}")
            return ""
    raise ValueError(f"不支持格式: {ext}")


# ──────────────────────────────────────────
# 2. 按标号解析（chunk 切割）
# ──────────────────────────────────────────


def parse_docx(path: str) -> list[dict]:
    """
    按标号切割：每个数字（1、2、3...）为一条 chunk。
    格式：标号 \\n 问题：... \\n 答复：...
    """
    text = _read_doc_text(path)
    if not text:
        return []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    entries = []
    entry_id = 1
    i = 0

    while i < len(lines):
        line = lines[i]
        # 标号：单独一行的纯数字
        if re.fullmatch(r"\d+", line):
            num = int(line)
            i += 1
            q_parts, a_parts = [], []
            while i < len(lines):
                cur = lines[i]
                if re.fullmatch(r"\d+", cur):
                    break  # 下一个标号
                if re.match(r"^DE\d+\s*疑难意见处理", cur):
                    i += 1
                    continue
                if cur.startswith("问题：") or cur.startswith("问题:"):
                    rest = re.sub(r"^问题[：:]", "", cur)
                    if "答复：" in rest or "答复:" in rest:
                        parts = re.split(r"答复[：:]", rest, maxsplit=1)
                        q_parts.append(parts[0].strip())
                        if parts[1].strip():
                            a_parts.append(parts[1].strip())
                    else:
                        q_parts.append(rest)
                    i += 1
                elif cur.startswith("答复：") or cur.startswith("答复:"):
                    a_parts.append(re.sub(r"^答复[：:]", "", cur))
                    i += 1
                else:
                    if a_parts:
                        a_parts.append(cur)
                    else:
                        q_parts.append(cur)
                    i += 1

            q = " ".join(q_parts).strip()
            a = " ".join(a_parts).strip()
            if q or a:
                m = re.search(r"(24-\d+|64-\d+|249\d+)", a)
                entries.append(
                    {
                        "id": entry_id,
                        "问题": q,
                        "答复": a,
                        "图号": m.group(1) if m else "",
                    }
                )
                entry_id += 1
        else:
            i += 1

    return entries


# ──────────────────────────────────────────
# 2. 构建索引（向量 + BM25）
# ──────────────────────────────────────────


def _tokenize_zh(text: str) -> list[str]:
    """中文分词（jieba），用于 BM25"""
    try:
        import jieba

        return list(jieba.cut(text))
    except ImportError:
        # jieba 未安装时按字符分词
        return list(text)


def build_index(entries: list[dict], force: bool = False):
    """向量化存入 ChromaDB，同时构建 BM25 索引"""
    from rank_bm25 import BM25Okapi

    collection = _get_collection()
    embedder = _get_embedder()

    existing = set(collection.get()["ids"])
    to_add = [e for e in entries if str(e["id"]) not in existing]

    if force:
        # 强制重建：清空后重新添加所有
        to_add = entries
        existing = set()

    if to_add:
        print(f"  向量化 {len(to_add)} 条新数据...")
        texts = [f"{e['问题']} {e['答复']}" for e in to_add]
        embeddings = embedder.encode(
            texts, batch_size=32, show_progress_bar=True
        ).tolist()

        if force and existing:
            collection.delete(ids=list(existing))

        collection.add(
            ids=[str(e["id"]) for e in to_add],
            embeddings=embeddings,
            documents=texts,
            metadatas=[
                {"问题": e["问题"], "答复": e["答复"], "图号": e.get("图号", "")}
                for e in to_add
            ],
        )
        print(f"  向量索引完成，共 {collection.count()} 条")

    # 重建 BM25（每次都重建保证与向量库一致）
    print("  构建 BM25 索引...")
    all_entries = entries  # 用传入的全量数据
    bm25_texts = [f"{e['问题']} {e['答复']}" for e in all_entries]
    bm25_metas = [
        {"问题": e["问题"], "答复": e["答复"], "图号": e.get("图号", "")}
        for e in all_entries
    ]
    tokenized = [_tokenize_zh(t) for t in bm25_texts]
    bm25_index = BM25Okapi(tokenized)

    os.makedirs(os.path.dirname(BM25_PATH), exist_ok=True)
    with open(BM25_PATH, "wb") as f:
        pickle.dump({"bm25": bm25_index, "docs": bm25_texts, "metas": bm25_metas}, f)

    # 缓存 entries 供后续加载
    with open(ENTRIES_PATH, "w", encoding="utf-8") as f:
        json.dump(all_entries, f, ensure_ascii=False, indent=2)

    # 刷新内存单例
    global _bm25, _bm25_docs, _bm25_metas
    _bm25 = bm25_index
    _bm25_docs = bm25_texts
    _bm25_metas = bm25_metas

    print("  BM25 索引完成")


# ──────────────────────────────────────────
# 3. 强化检索
# ──────────────────────────────────────────


def _rewrite_query(query: str) -> str:
    """
    查询改写：把口语化问题扩展为专业术语。
    例："VECS管子要直" → "VECS drain pipe 排放管 直线 净空 clearance upper deck"
    用本地 Qwen 做改写，失败则回退原始 query。
    """
    try:
        import model_loader

        messages = [
            {
                "role": "user",
                "content": (
                    f"你是船舶工程专家。请将以下问题改写为包含更多专业术语的搜索查询，"
                    f"用空格分隔关键词，只输出关键词，不超过20个词：\n{query}"
                ),
            }
        ]
        rewritten = model_loader.chat(messages, max_new_tokens=64).strip()
        # 防止模型输出过长或无关内容
        if rewritten and len(rewritten) < 200:
            return f"{query} {rewritten}"
    except Exception:
        pass
    return query


def _vector_search(query: str, top_k: int) -> list[dict]:
    """向量检索"""
    collection = _get_collection()
    embedder = _get_embedder()
    if collection.count() == 0:
        return []
    vec = embedder.encode([query]).tolist()
    results = collection.query(
        query_embeddings=vec,
        n_results=min(top_k, collection.count()),
        include=["metadatas", "distances"],
    )
    hits = []
    for meta, dist in zip(results["metadatas"][0], results["distances"][0]):
        hits.append(
            {
                "问题": meta["问题"],
                "答复": meta["答复"],
                "图号": meta.get("图号", ""),
                "score": float(1 - dist),  # cosine distance → similarity
                "via": "vector",
            }
        )
    return hits


def _bm25_search(query: str, top_k: int) -> list[dict]:
    """BM25 关键词检索"""
    bm25 = _get_bm25()
    if bm25 is None:
        return []
    tokens = _tokenize_zh(query)
    scores = bm25.get_scores(tokens)
    # 取 top_k
    top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[
        :top_k
    ]
    hits = []
    for idx in top_indices:
        if scores[idx] <= 0:
            continue
        hits.append(
            {
                "问题": _bm25_metas[idx]["问题"],
                "答复": _bm25_metas[idx]["答复"],
                "图号": _bm25_metas[idx].get("图号", ""),
                "score": float(scores[idx]),
                "via": "bm25",
            }
        )
    return hits


def _rerank(query: str, candidates: list[dict]) -> list[dict]:
    """用 cross-encoder 对候选结果重排序"""
    reranker = _get_reranker()
    if reranker is None or not candidates:
        return candidates
    pairs = [(query, f"{c['问题']} {c['答复']}") for c in candidates]
    scores = reranker.predict(pairs)
    for c, s in zip(candidates, scores):
        c["rerank_score"] = float(s)
    return sorted(candidates, key=lambda x: x["rerank_score"], reverse=True)


def _dedup(hits: list[dict]) -> list[dict]:
    """按问题文本去重"""
    seen, result = set(), []
    for h in hits:
        key = h["问题"][:50]
        if key not in seen:
            seen.add(key)
            result.append(h)
    return result


def search(query: str, top_k: int = 3, rewrite: bool = True) -> list[dict]:
    """
    强化混合检索：
      1. 查询改写（可选）
      2. 向量召回 top_k*2
      3. BM25 召回 top_k*2
      4. 合并去重
      5. Cross-encoder 重排序
      6. 相似度过滤 + 取 top_k

    返回：[{"问题": ..., "答复": ..., "图号": ..., "score": 0.xx}, ...]
    """
    # 1. 查询改写
    expanded_query = _rewrite_query(query) if rewrite else query

    # 2+3. 双路召回
    fetch_k = top_k * 2
    vec_hits = _vector_search(expanded_query, fetch_k)
    bm25_hits = _bm25_search(expanded_query, fetch_k)

    # 4. 合并去重（向量结果优先）
    candidates = _dedup(vec_hits + bm25_hits)

    # 5. 重排序
    candidates = _rerank(query, candidates)  # 重排用原始 query，不用扩展版

    # 6. 过滤 + 截断
    # 有重排分数时用重排分数过滤，否则用向量相似度过滤
    def _score(c):
        return c.get("rerank_score", c.get("score", 0))

    candidates = [c for c in candidates if _score(c) > SIMILARITY_THRESHOLD]
    results = candidates[:top_k]

    # 整理输出
    return [
        {
            "问题": r["问题"],
            "答复": r["答复"],
            "图号": r.get("图号", ""),
            "相似度": round(_score(r), 3),
            "来源": r.get("via", ""),
        }
        for r in results
    ]


# ──────────────────────────────────────────
# 4. 一键初始化
# ──────────────────────────────────────────


def _save_translated_docx(entries: list[dict], save_path: str):
    """把翻译后的条目写成结构化 docx，保存到 rag_docs/"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    title = doc.add_heading("疑难意见处理（已翻译整理）", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for e in entries:
        # 问题段落（加粗）
        p_q = doc.add_paragraph()
        run_id = p_q.add_run(f"[{e['id']}] 问题：")
        run_id.bold = True
        run_id.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p_q.add_run(e["问题"])

        # 答复段落
        p_a = doc.add_paragraph()
        run_a = p_a.add_run("答复：")
        run_a.bold = True
        run_a.font.color.rgb = RGBColor(0x37, 0x86, 0x3D)
        p_a.add_run(e["答复"])

        # 图号（如果有）
        if e.get("图号"):
            p_d = doc.add_paragraph()
            run_d = p_d.add_run("图号：")
            run_d.bold = True
            p_d.add_run(e["图号"])

        doc.add_paragraph()  # 条目间空行

    doc.save(save_path)
    print(f"  已保存整理后的 docx：{save_path}")


def init_rag_from_json(json_path: str):
    """从 *_translated.json 直接加载条目并建索引（无 doc 时使用）"""
    with open(json_path, "r", encoding="utf-8") as f:
        entries = json.load(f)
    if not entries:
        raise ValueError(f"JSON 为空：{json_path}")
    # 确保格式正确
    for i, e in enumerate(entries):
        if "问题" not in e or "答复" not in e:
            raise ValueError(f"JSON 格式错误，需含 问题/答复：{json_path}")
        e.setdefault("图号", "")
        e.setdefault("id", i + 1)
    print(f"  从 JSON 加载 {len(entries)} 条")
    print("  构建向量 + BM25 索引...")
    build_index(entries)
    print("  RAG 初始化完成")
    return entries


def init_rag(docx_path: str, json_path: str = None, force_retranslate: bool = False):
    """
    完整初始化：解析 → 翻译（中文版或缓存跳过）→ 建向量+BM25索引
    - 若为 *_中文版.docx，直接解析不翻译
    - 若有 *_translated.json 缓存，直接加载不翻译
    """
    if json_path is None:
        json_path = re.sub(r"\.(docx?)$", r"_translated.json", docx_path, flags=re.I)

    is_chinese_edition = "中文版" in os.path.basename(docx_path)
    entries = parse_docx(docx_path)
    print(f"  解析到 {len(entries)} 条（按标号切割）")

    skip_translate = is_chinese_edition or (
        entries and _is_chinese(entries[0]["问题"]) and _is_chinese(entries[0]["答复"])
    )
    if skip_translate and not force_retranslate:
        if not is_chinese_edition:
            print("  内容已是中文，跳过翻译...")
    elif os.path.exists(json_path) and not force_retranslate:
        print(f"  读取翻译缓存：{json_path}")
        with open(json_path, "r", encoding="utf-8") as f:
            entries = json.load(f)
    elif not skip_translate or force_retranslate:
        print("  开始翻译...")
        entries = _translate_all(entries)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(entries, f, ensure_ascii=False, indent=2)
        print(f"  翻译缓存已保存：{json_path}")

        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        translated_docx = os.path.join(
            os.path.dirname(docx_path), f"{base_name}_整理版.docx"
        )
        try:
            _save_translated_docx(entries, translated_docx)
        except Exception as e:
            print(f"  保存整理版 docx 失败：{e}")

    print("  构建向量 + BM25 索引...")
    build_index(entries)
    print("  RAG 初始化完成")
    return entries


def _translate_all(entries: list[dict]) -> list[dict]:
    import model_loader

    model_loader.load_model()
    total = len(entries)
    for i, e in enumerate(entries):
        print(f"  翻译 {i+1}/{total}", end="\r")
        e["问题"] = _translate(e["问题"], model_loader)
        e["答复"] = _translate(e["答复"], model_loader)
    print(f"\n  翻译完成")
    return entries


def _is_chinese(text: str) -> bool:
    """判断文本是否主要为中文"""
    if not text:
        return True
    cn = len(re.findall(r"[\u4e00-\u9fff]", text))
    return cn / max(len(text), 1) > 0.5


def _translate(text: str, ml) -> str:
    if not text:
        return text
    cn_ratio = len(re.findall(r"[\u4e00-\u9fff]", text)) / max(len(text), 1)
    if cn_ratio > 0.6:
        return text
    messages = [
        {
            "role": "user",
            "content": f"请将以下船舶工程文本翻译为简体中文，保留专业术语和图号，只输出翻译结果：\n{text}",
        }
    ]
    try:
        return ml.chat(messages, max_new_tokens=256).strip() or text
    except Exception:
        return text


# ──────────────────────────────────────────
# 命令行测试
# ──────────────────────────────────────────

if __name__ == "__main__":
    DOCX = os.path.join(os.path.dirname(os.path.dirname(__file__)), "rag_docs", "DE146疑难意见处理.docx")
    init_rag(DOCX)

    print("\nRAG 强化检索测试（输入 quit 退出）")
    while True:
        q = input("\n问题：").strip()
        if q.lower() == "quit":
            break
        hits = search(q, top_k=3)
        if not hits:
            print("  未找到相关内容")
            continue
        for i, h in enumerate(hits, 1):
            print(f"\n[{i}] 相似度={h['相似度']} 来源={h['来源']}")
            print(f"    问题：{h['问题'][:80]}")
            print(f"    答复：{h['答复'][:80]}")
            if h["图号"]:
                print(f"    图号：{h['图号']}")
